"""Generate PM2 Minutes of Meeting DOCX from meeting data."""

import copy
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "18.I.PM2-Template.v3.Minutes_of_Meeting.ProjectName.dd-mm-yyyy.vx_.x.docx",
)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "data", "mom")


def _set_cell(cell, text, bold=False):
    """Set cell text, preserving existing formatting."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(10)


def _set_cell_multiline(cell, text, bold=False):
    """Set cell text with line break support using separate paragraphs."""
    for p in cell.paragraphs:
        p.clear()
    lines = str(text).split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        if line.strip():
            run = p.add_run(line)
            run.bold = bold
            run.font.size = Pt(10)


def _add_shaded_row(table, cells_text, fill="E6E6E6", bold=True):
    """Add a header-style row with shading."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        if i < len(row.cells):
            _set_cell(row.cells[i], text, bold=bold)
            # Apply shading
            tc = row.cells[i]._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = tc.makeelement(qn("w:tcPr"), {})
                tc.insert(0, tcPr)
            shd = tcPr.find(qn("w:shd"))
            if shd is None:
                shd = tcPr.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): fill,
                    },
                )
                tcPr.append(shd)
            else:
                shd.set(qn("w:fill"), fill)


def generate_mom(meeting: dict) -> str:
    """Generate a PM2 MoM DOCX from meeting data. Returns output file path."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document(TEMPLATE_PATH)
    tables = doc.tables

    meeting_date = datetime.fromisoformat(meeting["created_at"])
    date_str = meeting_date.strftime("%d/%m/%Y")
    time_str = meeting_date.strftime("%H:%M")
    summary = meeting.get("summary") or {}
    segments = meeting.get("segments", [])

    # -- Cover page (Table 1) --
    cover = tables[1]
    _set_cell(cover.rows[3].cells[0], "Minutes of Meeting", bold=True)
    _set_cell(cover.rows[6].cells[0], f"Date: {date_str}")
    _set_cell(cover.rows[7].cells[0], "Doc. Version: 1.0")

    # -- Meeting info (Table 2) --
    info = tables[2]
    _set_cell(info.rows[0].cells[1], meeting.get("name", ""))
    _set_cell(info.rows[0].cells[3], f"{date_str} {time_str}")
    _set_cell(info.rows[1].cells[1], "Project Meeting")
    _set_cell(info.rows[1].cells[3], "Virtual")
    _set_cell(info.rows[2].cells[3], date_str)

    # -- Attendees (Table 3) — extract unique speakers --
    speakers = []
    seen = set()
    for seg in segments:
        spk = seg.get("speaker", "Unknown")
        if spk not in seen:
            seen.add(spk)
            speakers.append(spk)

    attendees = tables[3]
    # Fill first two template rows, then add more if needed
    for idx, spk in enumerate(speakers):
        if idx < 2:
            row = attendees.rows[idx + 1]
        else:
            row = attendees.add_row()
        _set_cell(row.cells[0], spk)
        initials = "".join(w[0].upper() for w in spk.split() if w) if " " in spk else spk[:2].upper()
        _set_cell(row.cells[1], initials)
        _set_cell(row.cells[2], "Yes")
        _set_cell(row.cells[3], "")

    # -- Meeting Agenda (Table 4) — numbered points --
    agenda_table = tables[4]
    agenda_items = summary.get("agenda", [])
    if not agenda_items:
        _set_cell(agenda_table.rows[1].cells[0], "Meeting agenda as discussed.")
    else:
        for idx, item in enumerate(agenda_items):
            if idx == 0:
                row = agenda_table.rows[1]
            else:
                row = agenda_table.add_row()
            _set_cell(row.cells[0], f"{idx + 1}. {item}")

    # -- Meeting Summary (Table 5) — structured bullet points --
    summary_table = tables[5]
    summary_text = summary.get("summary", "No summary available.")
    _set_cell_multiline(summary_table.rows[1].cells[0], summary_text)

    # -- Decisions (Table 6) --
    decisions = summary.get("decisions", [])
    dec_table = tables[6]
    for idx, decision in enumerate(decisions):
        if idx < 4:
            row = dec_table.rows[idx + 2]
        else:
            row = dec_table.add_row()
        _set_cell(row.cells[0], f"D-{idx + 1:03d}")
        _set_cell(row.cells[1], decision)
        _set_cell(row.cells[2], date_str)
        _set_cell(row.cells[3], "")

    # Clear unused placeholder rows in decisions
    used_dec = max(len(decisions), 1)
    for r in range(used_dec + 2, len(dec_table.rows)):
        for cell in dec_table.rows[r].cells:
            _set_cell(cell, "")

    # -- Action Items (Table 7) --
    actions = summary.get("action_items", [])
    act_table = tables[7]
    for idx, action in enumerate(actions):
        if idx < 4:
            row = act_table.rows[idx + 2]
        else:
            row = act_table.add_row()
        _set_cell(row.cells[0], f"A-{idx + 1:03d}")
        _set_cell(row.cells[1], date_str)
        _set_cell(row.cells[2], action)
        _set_cell(row.cells[3], "Open")
        _set_cell(row.cells[4], "")
        _set_cell(row.cells[5], "")

    # Clear unused placeholder rows in actions
    used_act = max(len(actions), 1)
    for r in range(used_act + 2, len(act_table.rows)):
        for cell in act_table.rows[r].cells:
            _set_cell(cell, "")

    # -- Next meeting (Table 8) --
    next_table = tables[8]
    _set_cell(next_table.rows[1].cells[0], "")

    # -- Save --
    safe_name = "".join(c if c.isalnum() or c in " -_" else "" for c in meeting.get("name", "meeting"))
    filename = f"MoM.{safe_name}.{date_str.replace('/', '-')}.v1.0.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)

    return output_path
